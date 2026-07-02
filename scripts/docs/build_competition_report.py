from __future__ import annotations

# ruff: noqa: E402

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict, Iterable, Sequence


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from mineshark.evaluation.competition import evaluate_scenarios, format_percent, load_scenarios


TITLE = "MineShark：面向 Tor 加密匿名通信的流量风险线索识别与大模型辅助研判系统"

DEFAULT_NETCLR_SUMMARY: Dict[str, Any] = {
    "task_framing": "Tor encrypted anonymous-communication traffic risk-clue binary baseline",
    "binary_view": {
        "normal_dir": "datasets/experiments/ppi/tor/netclr_drift_binary/normal",
        "risk_dir": "datasets/experiments/ppi/tor/netclr_drift_binary/risk",
        "negative_label_name": "netclr_inferior_condition",
        "positive_label_name": "netclr_superior_condition",
        "label_warning": "These are NetCLR condition labels, not normal vs malware labels.",
    },
    "quality": {
        "file_count": 2,
        "sample_count": 28312,
        "invalid_rows": 0,
        "class_count": 93,
        "average_sequence_length": 127.47739474427804,
        "min_sequence_length": 22,
        "max_sequence_length": 128,
        "short_sample_count": 0,
        "empty_direction_sample_count": 0,
    },
    "checkpoint": "checkpoints/tor_netclr_drift_binary_gpu_v1.pt",
    "metrics": {
        "sample_count": 28312,
        "threshold": 0.5664353178573608,
        "accuracy": 0.29461005933879625,
        "precision": 0.8290482634190347,
        "recall": 0.08576761549230051,
        "f1": 0.15545312301771894,
        "fpr": 0.055071200232490555,
        "fnr": 0.9142323845076995,
        "tp": 1838,
        "fp": 379,
        "tn": 6503,
        "fn": 19592,
    },
    "interpretation": (
        "At a low-FPR operating point, precision is high but recall is very low. "
        "Report this as high-confidence traffic-side risk clues for assisted triage, "
        "not as a mature Tor threat detector."
    ),
}


def set_run_font(run, *, name: str = "宋体", size: float | None = 12, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_para(
    doc: Document,
    text: str = "",
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool | None = None,
    size: float = 12,
    style: str | None = None,
) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size={1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: Any, *, bold: bool = False, fill: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=10.5, bold=bold)


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, fill="F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def fmt_decimal(value: Any, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def add_cover(doc: Document) -> None:
    add_para(doc, "附件2：", size=12)
    add_para(
        doc,
        "第十九届全国大学生信息安全竞赛（作品赛）暨第三届“长城杯”网数智安全大赛（作品赛）",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=14,
    )
    add_para(doc, "作品报告", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22)
    add_para(doc, "■命题赛道             □自由赛道", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "", size=12)
    add_para(doc, f"作品名称：{TITLE}", size=12)
    add_para(doc, "电子邮箱：匿名提交，系统填报为准", size=12)
    add_para(doc, "提交日期：2026年7月", size=12)
    add_para(doc, "", size=12)
    add_para(doc, "匿名化说明", bold=True, size=12)
    add_para(
        doc,
        "本报告为网评匿名版，仅保留作品设计、实现、测试与分析内容，不包含任何可识别参赛身份或真实运行凭据的信息。",
        size=12,
    )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_para(doc, "目     录", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    for item in [
        "摘要",
        "第一章 作品概述",
        "第二章 作品设计与实现",
        "第三章 作品测试与分析",
        "第四章 创新性说明",
        "第五章 总结",
        "参考文献",
    ]:
        add_para(doc, item, size=12)
    doc.add_page_break()


def add_summary(doc: Document) -> None:
    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "MineShark 面向 Tor 加密匿名通信及其他加密通信协议中的流量风险线索识别问题，围绕“不解密明文也能提取可复核风险线索”的目标，"
        "构建了加密流量元数据分析、Transformer 风险线索判定、多源安全证据聚合和大模型辅助研判的原型系统。"
        "系统从 Tor/NetCLR 数据、MineShark/Zeek 风格日志、Wazuh 告警、Suricata 规则告警和本地安全知识库中抽取证据，"
        "以包长序列、方向序列、包间隔、端口和连接上下文作为主要输入，输出风险线索分、证据链、误报边界和人工复核建议。",
    )
    add_para(
        doc,
        "项目早期以通用加密恶意流量检测为目标，完成日志解析、特征转换、Transformer 二分类训练和检测结果输出；"
        "后续接入 Wazuh、Zeek、Suricata、RAG、Agent 和 Console，形成安全研判系统原型；随着研究方向聚焦到 Tor，"
        "当前最终实验主线采用 NetCLR Tor 条件漂移数据构建二分类流量风险线索基线，标签为 "
        "`netclr_inferior_condition` 与 `netclr_superior_condition`。该标签表示网络条件差异下的流量侧风险线索，"
        "不是恶意/正常事实。WFlib CW 单标签页链路仅作为 95 类 closed-world 网站指纹备用实验，不作为最终二分类主线。"
        "系统保留 LLM/RAG/Agent 作为辅助研判能力，用于把模型风险线索、Wazuh/Zeek/Suricata 证据和本地 playbook "
        "组织成可复核的中文报告。系统不自动封禁、不写回 Wazuh 状态，也不把模型概率直接等同于攻击事实。",
    )
    add_para(
        doc, "关键词：Tor 加密匿名通信；流量风险线索识别；NetCLR；Transformer；Wazuh；Zeek；Suricata；RAG；安全研判", size=12
    )


def add_overview(doc: Document) -> None:
    add_heading(doc, "第一章 作品概述", 1)
    add_heading(doc, "1.1 背景分析", 2)
    add_para(
        doc,
        "HTTPS、SSH、DNS over TLS 等加密通信已经成为企业网络的常态。传统依赖明文 payload、特征串或固定规则的检测方法"
        "在加密场景下受到限制，但连接元数据仍然保留了行为模式，例如包长、方向、连接持续时间、包间隔、端口、通信频率和"
        "同一主机的相邻告警。MineShark 利用这些元数据识别 C2 Beacon、异常隧道、SSH 暴力破解后操作、异常命令序列以及"
        "Tor 条件漂移场景中的流量侧风险线索。",
    )
    add_para(
        doc,
        "Tor 是匿名加密通信协议，本身不是攻击事实，Tor 用户也不等于恶意用户。因此本作品不声称“检测 Tor 恶意用户”，"
        "而是将 Tor/NetCLR 实验限定为流量风险线索二分类基线：模型输出用于辅助研判和优先级排序，最终结论仍需结合日志、规则、"
        "资产上下文和人工复核。",
    )
    add_heading(doc, "1.2 作品目标", 2)
    for item in [
        "在不解密 TLS/SSH/Tor 明文的前提下，完成加密流量元数据与风险线索的对比分析。",
        "以 NetCLR Tor 条件漂移数据构建二分类流量风险线索基线，并清晰说明其标签边界。",
        "通过 Transformer 风险线索分和阈值策略输出可解释、可复核的流量侧风险线索。",
        "关联 Wazuh、Zeek、Suricata 和本地 RAG playbook，生成可复核的中文研判报告。",
        "提供 live Wazuh/WSL 演示路径和离线 fallback 演示路径，降低比赛现场环境风险。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.3 应用前景", 2)
    add_para(
        doc,
        "本作品适用于校园网、企业内网和实验 SOC 场景中的加密流量巡检。它不替代现有 IDS/SIEM，而是作为旁路分析组件，"
        "把 AI 模型风险线索接入 Wazuh 告警体系，再由 Agent 将多源证据整理为人工可读的事件说明，降低人工研判成本。",
    )


def add_design(doc: Document) -> None:
    add_heading(doc, "第二章 作品设计与实现", 1)
    add_heading(doc, "2.1 系统总体方案", 2)
    add_para(doc, "图 1  系统总体流程（文本化表示）", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
    flow_rows = [
        ("输入层", "MineShark/Zeek 连接日志、Wazuh 告警、Suricata eve.json、本地安全 playbook"),
        ("检测层", "解析包长、方向、IAT、端口等元数据，使用 Transformer 输出风险线索分"),
        ("证据层", "按 alert_id、UID、IP 和时间窗口查询 Wazuh、Zeek、Suricata 与 RAG"),
        ("研判层", "EvidenceBundle、质量检查、LLM 或确定性报告生成"),
        ("展示层", "Markdown/JSON 报告、tool_trace、MineShark Console"),
    ]
    add_table(doc, ["层次", "功能"], flow_rows)
    add_heading(doc, "2.2 加密流量元数据检测", 2)
    add_para(
        doc,
        "检测模型不读取会话明文，而是以包长序列、方向序列、包间隔时间和连接上下文作为输入。"
        "Transformer 用于建模序列中不同位置之间的依赖关系，输出二分类风险线索分。"
        "在 NetCLR 主线中，负侧标签为 `netclr_inferior_condition`，正侧标签为 `netclr_superior_condition`；"
        "它们是网络条件差异下的风险线索标签，不是正常/恶意标签。训练分支保留阈值校准逻辑，优先控制误报率，"
        "而不是只追求单一准确率。",
    )
    add_heading(doc, "2.3 多源证据聚合", 2)
    add_table(
        doc,
        ["证据源", "作用", "边界"],
        [
            ("MineShark AI 告警", "提供模型风险分、UID、五元组和时间线索", "只能作为风险线索"),
            ("Wazuh", "提供平台告警和规则摄取结果", "API 不通时回退本地 alerts.json"),
            ("Zeek", "提供连接级元数据和 UID 上下文", "依赖日志采集稳定性"),
            ("Suricata", "提供 IDS 规则侧旁证", "规则命中不等同于入侵事实"),
            ("RAG playbook", "提供 C2、隧道、误报治理等研判知识", "离线时可降级读取 JSONL"),
        ],
    )
    add_heading(doc, "2.4 大模型辅助研判", 2)
    add_para(
        doc,
        "LLM 在本作品中不直接承担检测任务，而是读取结构化证据包并生成中文研判报告。JSON 报告保留 `tool_trace`，"
        "记录每次工具调用的参数和返回结果，避免报告成为不可追溯的黑盒输出。若外部大模型或 FAISS 索引不可用，系统仍可通过"
        " `--evidence-only` 和 JSONL playbook fallback 生成确定性报告。",
    )
    add_heading(doc, "2.5 安全边界", 2)
    for item in [
        "不自动封禁 IP、隔离主机或修改防火墙。",
        "不写回 Wazuh 告警状态，不替换现有 Wazuh、Zeek、Suricata 服务。",
        "不把模型概率作为唯一判定依据，最终结论需要人工复核。",
    ]:
        add_bullet(doc, item)


def add_netclr_experiment(doc: Document, summary: Dict[str, Any]) -> None:
    binary_view = summary["binary_view"]
    quality = summary["quality"]
    metric = summary["metrics"]
    add_heading(doc, "3.2 NetCLR Tor 二分类流量风险线索基线", 2)
    add_para(
        doc,
        "最终实验包采用 NetCLR Tor 条件漂移数据构建二分类流量风险线索基线。该实验不是 Tor 恶意用户检测，"
        "也不是 WFlib CW 的 95 类 closed-world 网站指纹分类。WFlib 单标签页链路仅作为备用工程能力保留。",
    )
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("负侧训练视图", binary_view["normal_dir"]),
            ("正侧训练视图", binary_view["risk_dir"]),
            ("负侧报告标签", binary_view["negative_label_name"]),
            ("正侧报告标签", binary_view["positive_label_name"]),
            ("Checkpoint", summary["checkpoint"]),
            ("标签边界", "NetCLR 条件差异风险线索，不是正常/恶意事实标签"),
        ],
    )
    add_table(
        doc,
        ["质量检查项", "数值"],
        [
            ("file_count", quality["file_count"]),
            ("sample_count", quality["sample_count"]),
            ("invalid_rows", quality["invalid_rows"]),
            ("class_count", quality["class_count"]),
            ("average_sequence_length", fmt_decimal(quality["average_sequence_length"])),
            ("min_sequence_length", quality["min_sequence_length"]),
            ("max_sequence_length", quality["max_sequence_length"]),
            ("short_sample_count", quality["short_sample_count"]),
            ("empty_direction_sample_count", quality["empty_direction_sample_count"]),
        ],
    )
    add_table(
        doc,
        ["指标", "数值"],
        [
            ("sample_count", metric["sample_count"]),
            ("threshold", fmt_decimal(metric["threshold"])),
            ("accuracy", format_percent(metric["accuracy"])),
            ("precision", format_percent(metric["precision"])),
            ("recall", format_percent(metric["recall"])),
            ("f1", format_percent(metric["f1"])),
            ("fpr", format_percent(metric["fpr"])),
            ("fnr", format_percent(metric["fnr"])),
            ("TP / FP / TN / FN", f"{metric['tp']} / {metric['fp']} / {metric['tn']} / {metric['fn']}"),
        ],
    )
    add_para(
        doc,
        "结果解读：在低误报约束下，NetCLR 二分类基线 precision 较高，但 recall 很低。这说明模型可以输出一部分"
        "高置信流量侧风险线索，但漏检严重。因此当前结果适合用于风险线索优先级排序和辅助研判，不适合包装成成熟的"
        "Tor 恶意检测系统。",
    )


def add_testing(doc: Document, metrics: Dict[str, Any], netclr_summary: Dict[str, Any]) -> None:
    add_heading(doc, "第三章 作品测试与分析", 1)
    add_heading(doc, "3.1 测试方案", 2)
    add_para(
        doc,
        "测试采用 NetCLR Tor 最终实验包、脱敏小型竞赛评测集与系统级演示三条路径。NetCLR 实验包用于支撑当前"
        "Tor 风险线索二分类主线；脱敏竞赛评测集用于展示普通加密通信、C2 Beacon、加密隧道和 SSH 异常行为的"
        "风险线索对比；系统级演示包括 live Wazuh/WSL 链路和离线 fallback 链路。",
    )
    add_netclr_experiment(doc, netclr_summary)
    add_heading(doc, "3.3 脱敏竞赛场景覆盖", 2)
    category_rows = []
    for item in metrics["categories"]:
        category_rows.append(
            [
                item["category"],
                item["count"],
                f"{item['average_score']:.3f}",
                format_percent(item["accuracy"]),
                format_percent(item["fpr"]),
            ]
        )
    add_table(doc, ["场景", "样本数", "平均风险分", "Accuracy", "FPR"], category_rows)
    add_heading(doc, "3.4 脱敏竞赛场景指标", 2)
    metric = metrics["metrics"]
    add_table(
        doc,
        ["指标", "数值"],
        [
            ("Accuracy", format_percent(metric["accuracy"])),
            ("Precision", format_percent(metric["precision"])),
            ("Recall", format_percent(metric["recall"])),
            ("F1", format_percent(metric["f1"])),
            ("FPR", format_percent(metric["fpr"])),
        ],
    )
    add_table(
        doc,
        ["混淆矩阵项", "数量"],
        [
            ("TP", metric["tp"]),
            ("FP", metric["fp"]),
            ("TN", metric["tn"]),
            ("FN", metric["fn"]),
        ],
    )
    add_heading(doc, "3.5 误报与漏报分析", 2)
    add_para(
        doc,
        "脱敏竞赛 fixture 的默认阈值为 0.70，出现 1 个误报和 1 个漏报。误报样例是自动化巡检 SSH 会话，"
        "固定长度往返和周期性行为与异常隧道存在相似性；漏报样例是低频长周期 C2 Beacon，风险分略低于阈值。"
        "这些样例用于说明模型概率只是风险线索，不等于攻击事实。后续优化方向是引入业务白名单、长周期统计窗口和"
        "人工反馈闭环。",
    )
    add_heading(doc, "3.6 演示验证", 2)
    add_table(
        doc,
        ["演示路径", "命令", "输出"],
        [
            (
                "竞赛评测",
                "python scripts/eval/run_competition_eval.py --scenario-dir tests/fixtures/competition_scenarios",
                "metrics.json、comparison.md、table_data.csv",
            ),
            (
                "离线 fallback",
                "python scripts/agent/run_offline_fixture_demo.py --fixture-dir tests/fixtures/demo_event",
                "offline_agent_report.json、offline_agent_report.md",
            ),
            (
                "Live Wazuh/WSL",
                "powershell -ExecutionPolicy Bypass -File scripts/agent/run_wsl_cli_agent_demo.ps1",
                "agent_audit_report.json、agent_audit_report.md、Console 展示",
            ),
        ],
    )


def add_innovation(doc: Document) -> None:
    add_heading(doc, "第四章 创新性说明", 1)
    innovations = [
        ("不解密条件下的流量风险线索识别", "使用包长、方向、IAT、端口和连接上下文识别加密通信中的可疑行为模式。"),
        (
            "Tor/NetCLR 任务边界治理",
            "明确 Tor 不是攻击事实，NetCLR 二分类标签是条件漂移风险证据，不包装成恶意用户检测。",
        ),
        ("旁路式安全架构", "不替换现有 Wazuh/Zeek/Suricata，只读取告警和日志，降低接入风险。"),
        ("误报治理导向", "训练分支保留阈值校准、良性对照和人工复核模板，面向实际安全运营。"),
        ("可追溯 Agent 报告", "JSON 中保留 evidence_bundle、quality_checks 和 tool_trace，便于复核。"),
        ("离线可复现演示", "外部大模型、FAISS 或 Wazuh API 不可用时，仍可用 fixture 生成同结构报告。"),
    ]
    add_table(doc, ["创新点", "说明"], innovations)


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "第五章 总结", 1)
    add_para(
        doc,
        "MineShark 围绕第七题要求实现了加密通信流量分析、流量风险线索识别模型和正常/风险流量对比实验，并通过 Wazuh、Zeek、"
        "Suricata、RAG 与 Agent 把模型风险线索转化为可复核报告。当前最终实验主线已经收束为 NetCLR Tor 二分类流量风险线索基线；"
        "其低误报约束下 precision 较高，但 recall 很低，适合用于高置信风险线索提示和辅助研判，不适合包装成成熟检测系统。"
        "后续可继续扩展更合理的标签定义、更丰富的流量特征、更多协议类型、业务白名单、前端可视化和报告质量自动评估。",
    )
    add_para(
        doc,
        "作品的工程边界保持克制：不自动处置、不把 Tor 流量等同于攻击事实、不夸大模型结论、不提交真实密钥或大规模原始流量数据。"
        "这样的设计更符合安全系统渐进式接入和人工复核的实际要求。",
    )


def add_references(doc: Document) -> None:
    add_heading(doc, "参考文献", 1)
    refs = [
        "Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. NeurIPS, 2017.",
        "Paxson V. Bro: A System for Detecting Network Intruders in Real-Time[J]. Computer Networks, 1999.",
        "Anderson B, McGrew D. Machine Learning for Encrypted Malware Traffic Classification[C]. KDD, 2016.",
        "NIST. Guide to Intrusion Detection and Prevention Systems (IDPS): SP 800-94[S].",
        "Zeek Project. Zeek Documentation[EB/OL].",
        "Suricata Project. Suricata User Guide[EB/OL].",
        "Wazuh Project. Wazuh Documentation[EB/OL].",
        "MineShark 项目文档：README、competition_submission、tor_dataset_strategy、final_tor_netclr_experiment_package、reporting。",
    ]
    for ref in refs:
        add_para(doc, ref, size=12)


def load_metrics(path: Path) -> Dict[str, Any]:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    scenarios = load_scenarios(ROOT / "tests" / "fixtures" / "competition_scenarios")
    return evaluate_scenarios(scenarios, threshold=0.70)


def load_netclr_summary(path: Path) -> Dict[str, Any]:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return DEFAULT_NETCLR_SUMMARY


def build_report(output: Path, metrics_path: Path, netclr_summary_path: Path) -> None:
    metrics = load_metrics(metrics_path)
    netclr_summary = load_netclr_summary(netclr_summary_path)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)
    add_summary(doc)
    add_overview(doc)
    add_design(doc)
    add_testing(doc, metrics, netclr_summary)
    add_innovation(doc)
    add_conclusion(doc)
    add_references(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Build the anonymous MineShark competition report DOCX.")
    parser.add_argument("--metrics-json", default="outputs/competition/metrics.json")
    parser.add_argument("--netclr-summary-json", default="outputs/final_tor_netclr_package/metrics_summary.json")
    parser.add_argument("--output", default="outputs/submission/MineShark_第七题作品报告_匿名版.docx")
    return parser


def main() -> None:
    args = build_parser().parse_args()
    output = Path(args.output)
    if not output.is_absolute():
        output = ROOT / output
    metrics_path = Path(args.metrics_json)
    if not metrics_path.is_absolute():
        metrics_path = ROOT / metrics_path
    netclr_summary_path = Path(args.netclr_summary_json)
    if not netclr_summary_path.is_absolute():
        netclr_summary_path = ROOT / netclr_summary_path
    build_report(output, metrics_path, netclr_summary_path)
    print(f"DOCX report: {output}")


if __name__ == "__main__":
    main()
